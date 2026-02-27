import sys
from docx import Document
from docx.shared import Pt

def create_template(input_path, output_path):
    print("Loading", input_path)
    doc = Document(input_path)
    
    # 1. Update the header paragraph
    # It contains "鐘 點 簽 名 表 (Class)"
    for p in doc.paragraphs:
        if "鐘 點 簽 名 表" in p.text:
            # We want to replace just the (Class) part or the whole text to be safe
            p.text = "鐘 點 簽 名 表 ({className} -- {name} - {monthName} {year})"
            # Restore basic formatting if we overwrote it
            for run in p.runs:
                run.font.size = Pt(28) # Assuming around 28pt based on heading
    
    # 2. Update the Table
    table = doc.tables[0]
    
    # Row 0: Header (keep as is)
    # Row 1-23: Data Rows
    # Row 24: Signature (keep as is)
    
    # We want to replace row 1 with the loop start, row 2 with data, row 3 with loop end
    # But for docxtemplater, table loops are easiest done by wrapping a single row in {#days} ... {/days}
    # However, docxtemplater supports {#days} inside the first cell and {/days} in the last cell of the same row!
    
    # Let's clear out rows 2 to 23
    for _ in range(22):
        # We delete row 2 repeatedly (which shifts the others up)
        tbl_ele = table._element
        row_ele = table.rows[2]._tr
        tbl_ele.remove(row_ele)
        
    # Now table has: Header (Row 0), One Data Row (Row 1), Signature (Row 2)
    data_row = table.rows[1]
    
    # Cells in data row: # | Date | Class | Time | Minutes | Remarks
    data_row.cells[0].text = "{#days}{rowIndex}"
    data_row.cells[1].text = "{date}"
    data_row.cells[2].text = "{className}"
    data_row.cells[3].text = "{time}"
    data_row.cells[4].text = "{mins}"
    data_row.cells[5].text = "{remarks}{/days}"
    
    # Ensure fonts are somewhat maintained
    for cell in data_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(14) # Approx 14pt (28 half-points)
                
    doc.save(output_path)
    print("Saved template to", output_path)

if __name__ == "__main__":
    create_template('timesheet/SignatureSheet.docx', 'timesheet/template.docx')
